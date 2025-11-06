"""
Text extraction module for AI Resume Analyzer.

This module handles extracting text from PDF and DOCX files,
with Nutrient OCR fallback for complex documents.
"""

import os
import logging
from typing import Optional

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_text_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pdfplumber.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as string
        
    Raises:
        Exception: If PDF extraction fails
    """
    try:
        import pdfplumber
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            if not pdf.pages:
                raise ValueError("PDF file contains no pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean up whitespace while preserving structure
                        cleaned_text = '\n'.join(line.strip() for line in page_text.split('\n') if line.strip())
                        text_content.append(cleaned_text)
                    else:
                        logger.warning(f"No text extracted from page {page_num}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    continue
        
        if not text_content:
            raise ValueError("No text could be extracted from the PDF")
        
        # Join pages with double newline for clear separation
        full_text = '\n\n'.join(text_content)
        
        # Additional cleanup for common PDF artifacts
        full_text = full_text.replace('\x00', '')  # Remove null characters
        full_text = full_text.replace('\ufeff', '')  # Remove BOM
        
        logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
        return full_text
        
    except ImportError:
        raise ImportError("pdfplumber library is required for PDF extraction. Install with: pip install pdfplumber")
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as string
        
    Raises:
        Exception: If DOCX extraction fails
    """
    try:
        from docx import Document
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Invalid DOCX file format: {str(e)}")
        
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        if not text_content:
            raise ValueError("No text could be extracted from the DOCX file")
        
        # Join paragraphs with single newline to preserve structure
        full_text = '\n'.join(text_content)
        
        # Clean up common artifacts
        full_text = full_text.replace('\x00', '')  # Remove null characters
        full_text = full_text.replace('\ufeff', '')  # Remove BOM
        full_text = full_text.replace('\r\n', '\n')  # Normalize line endings
        full_text = full_text.replace('\r', '\n')
        
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except ImportError:
        raise ImportError("python-docx library is required for DOCX extraction. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def needs_nutrient_ocr(text: str) -> bool:
    """
    Assess if extracted text quality is poor and needs Nutrient OCR.
    
    Args:
        text: Extracted text to assess
        
    Returns:
        True if Nutrient OCR is needed, False otherwise
    """
    if not text or len(text.strip()) < 50:
        logger.info("Text too short, recommending Nutrient OCR")
        return True
    
    # Calculate character density metrics
    total_chars = len(text)
    whitespace_chars = len([c for c in text if c.isspace()])
    non_whitespace_chars = total_chars - whitespace_chars
    
    if non_whitespace_chars == 0:
        logger.info("No non-whitespace characters found, recommending Nutrient OCR")
        return True
    
    # Check for excessive whitespace (indicates formatting issues)
    whitespace_ratio = whitespace_chars / total_chars
    if whitespace_ratio > 0.7:
        logger.info(f"High whitespace ratio ({whitespace_ratio:.2f}), recommending Nutrient OCR")
        return True
    
    # Check for multi-column layout indicators
    lines = text.split('\n')
    short_lines = [line for line in lines if len(line.strip()) > 0 and len(line.strip()) < 20]
    if len(short_lines) > len(lines) * 0.6:
        logger.info("Many short lines detected (possible multi-column), recommending Nutrient OCR")
        return True
    
    # Check for formatting artifacts
    artifact_patterns = [
        '...',  # Excessive dots
        '___',  # Excessive underscores
        '   ',  # Multiple spaces
        '\n\n\n',  # Excessive newlines
    ]
    
    artifact_count = sum(text.count(pattern) for pattern in artifact_patterns)
    if artifact_count > 10:
        logger.info(f"High formatting artifact count ({artifact_count}), recommending Nutrient OCR")
        return True
    
    # Check character density per line
    valid_lines = [line.strip() for line in lines if line.strip()]
    if valid_lines:
        avg_line_length = sum(len(line) for line in valid_lines) / len(valid_lines)
        if avg_line_length < 15:  # Very short average line length
            logger.info(f"Low average line length ({avg_line_length:.1f}), recommending Nutrient OCR")
            return True
    
    # Check for readable text patterns
    import re
    word_pattern = re.compile(r'\b[a-zA-Z]{2,}\b')
    words = word_pattern.findall(text)
    
    if len(words) < 20:  # Too few recognizable words
        logger.info(f"Too few recognizable words ({len(words)}), recommending Nutrient OCR")
        return True
    
    # Calculate readability score based on word/character ratio
    word_chars = sum(len(word) for word in words)
    readability_ratio = word_chars / non_whitespace_chars if non_whitespace_chars > 0 else 0
    
    if readability_ratio < 0.5:  # Less than 50% of non-whitespace chars are in words
        logger.info(f"Low readability ratio ({readability_ratio:.2f}), recommending Nutrient OCR")
        return True
    
    logger.info("Text quality appears good, no need for Nutrient OCR")
    return False


def call_nutrient_ocr(file_path: str, api_key: str) -> str:
    """
    Call Nutrient OCR service for text extraction fallback.
    
    Args:
        file_path: Path to the file for OCR
        api_key: Nutrient API key
        
    Returns:
        OCR extracted text as string
        
    Raises:
        Exception: If Nutrient API call fails
    """
    try:
        import requests
        import time
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not api_key:
            raise ValueError("Nutrient API key is required")
        
        # Prepare API request
        url = "https://api.nutrient.io/build"
        
        headers = {
            "Authorization": f"Bearer {api_key}"
        }
        
        # Prepare multipart form data
        try:
            with open(file_path, 'rb') as f:
                files = {
                    'file': (os.path.basename(file_path), f, 'application/octet-stream')
                }
                
                # Instructions for OCR text extraction
                instructions = {
                    "parts": [
                        {
                            "file": "file",
                            "pages": {
                                "start": 0,
                                "end": -1
                            }
                        }
                    ],
                    "actions": [
                        {
                            "type": "ocr",
                            "language": "english"
                        },
                        {
                            "type": "export",
                            "format": "txt"
                        }
                    ]
                }
                
                data = {
                    'instructions': str(instructions).replace("'", '"')
                }
                
                # Make API request with retry logic
                max_retries = 2
                retry_delay = 1
                
                for attempt in range(max_retries + 1):
                    try:
                        logger.info(f"Calling Nutrient OCR API (attempt {attempt + 1})")
                        response = requests.post(url, headers=headers, files=files, data=data, timeout=60)
                        
                        if response.status_code == 200:
                            # Nutrient returns the processed text directly
                            extracted_text = response.text.strip()
                            
                            if not extracted_text:
                                raise ValueError("Nutrient returned empty text")
                            
                            logger.info(f"Successfully extracted {len(extracted_text)} characters via Nutrient OCR")
                            return extracted_text
                        
                        elif response.status_code == 429:  # Rate limit
                            if attempt < max_retries:
                                wait_time = retry_delay * (2 ** attempt)
                                logger.warning(f"Rate limited, waiting {wait_time} seconds before retry")
                                time.sleep(wait_time)
                                continue
                            else:
                                raise Exception("Rate limit exceeded, max retries reached")
                        
                        elif response.status_code == 401:
                            raise Exception("Invalid Nutrient API key")
                        
                        elif response.status_code == 400:
                            try:
                                error_data = response.json()
                                error_msg = error_data.get('error', 'Bad request')
                            except:
                                error_msg = response.text
                            raise Exception(f"Nutrient API error: {error_msg}")
                        
                        else:
                            raise Exception(f"Nutrient API returned status {response.status_code}: {response.text}")
                            
                    except requests.exceptions.Timeout:
                        if attempt < max_retries:
                            logger.warning(f"Request timeout, retrying in {retry_delay} seconds")
                            time.sleep(retry_delay)
                            continue
                        else:
                            raise Exception("Nutrient API request timed out after retries")
                            
                    except requests.exceptions.ConnectionError:
                        if attempt < max_retries:
                            logger.warning(f"Connection error, retrying in {retry_delay} seconds")
                            time.sleep(retry_delay)
                            continue
                        else:
                            raise Exception("Failed to connect to Nutrient API after retries")
                
                raise Exception("Nutrient OCR failed after all retry attempts")
                
        except Exception as e:
            if "File not found" in str(e) or "API key" in str(e):
                raise
            raise Exception(f"Failed to process file for Nutrient OCR: {str(e)}")
        
    except ImportError:
        raise ImportError("requests library is required for Nutrient integration. Install with: pip install requests")
    except Exception as e:
        logger.error(f"Nutrient OCR failed for {file_path}: {str(e)}")
        raise Exception(f"Nutrient OCR extraction failed: {str(e)}")